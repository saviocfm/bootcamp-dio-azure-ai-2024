import requests
import os
from docx import Document
from bs4 import BeautifulSoup

sub_key =  ''
endpoint = "https://api.cognitive.microsofttranslator.com"
location = "eastus2"
language = 'pt-br'

params = {
    'api-version':'3.0',
    'from':'en',
    'to':[language]
}

def translator_text(text,target_language=language):
    path = '/translate'
    constructed_url = endpoint+path
    headers={
        'Ocp-Apim-Subscription-Key':sub_key,
        'Ocp-Apim-Subscription-Region':location,
        'Content-type':'application/json',
        'C-ClientTraceId':str(os.urandom(16))
    }
    
    body = [
        {
            'text':text
        }
    ]
    request = requests.post(constructed_url,params=params,headers=headers,json=body)
    res = request.json()
    return res[0]["translations"][0]["text"]


def translate_doc(path):
    document = Document(path)
    fulltext = []
    for paragraph in document.paragraphs:
        fulltext.append(translator_text(paragraph))

    translated_doc = Document()
    for linha in fulltext:
        translate_doc.add_paragraph(linha)

    path_translated = path.replace(".docx",f"_{language}_translated.docx")

    translate_doc.save(path_translated)
    return path_translated

def extract_text_from_url(url):
    response = requests.get(url)

    if response.status_code != 200:
        return None
    else: 
        soup = BeautifulSoup(response.text,'html.parser')
        for script_or_style in soup(["script","style"]):
            script_or_style.decompose()
        texto = soup.get_text(separator=' ')
        linhas = (line.strip() for line in texto.splitlines())
        parts = (phrase.strip() for line in linhas for phrase in line.split("   "))
        texto_limpo = '\n'.join(line for line in linhas if line)
        return texto_limpo
    
from langchain_openai.chat_models.azure import AzureChatOpenAI

client = AzureChatOpenAI(
    azure_endpoint=endpoint,
    api_version='2024-02-15-preview',
    api_key=sub_key,
    deployment_name="gpt-4o-mini",
    max_retries=0
)

def translate_article(text,lang=language):
    messages = [
        ("system","você atua como tradutor de textos"),
        ("user",f"traduza o {text} para o idioma {lang} e responda em markdown")
    ]

    response = client.invoke(messages)
    print(response.content)
    return response.content